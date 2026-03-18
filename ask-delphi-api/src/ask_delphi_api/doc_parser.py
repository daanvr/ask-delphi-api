"""
Document parser: Word → JSON conversie.
"""
from pathlib import Path
import numpy as np
import pandas as pd
import warnings
import re
import json
import tempfile
from docx.oxml.ns import qn

from ask_delphi_api import api


def read_dir(dir_path):
    dir = Path(dir_path)
    if not dir.is_dir():
        raise ValueError("taaksjabloon directory not found")

    paths = []
    for path in dir.iterdir():
        if path.is_file():
            print(path.name)
            paths.append(path.absolute())

    return paths


def extract_image_from_run(run, document):
    """Haal image bytes en extensie op uit een run met een inline image."""
    drawing_els = run._r.findall(qn("w:drawing"))
    if not drawing_els:
        return None

    for drawing in drawing_els:
        blips = drawing.findall(f".//{qn('a:blip')}")
        for blip in blips:
            rId = blip.get(qn("r:embed"))
            if rId is None:
                continue
            rel = run.part.rels.get(rId)
            if rel is None:
                continue
            image_part = rel.target_part
            content_type = image_part.content_type
            image_bytes = image_part.blob

            ext_map = {
                "image/png": ".png",
                "image/jpeg": ".jpg",
                "image/gif": ".gif",
                "image/bmp": ".bmp",
                "image/tiff": ".tiff",
                "image/svg+xml": ".svg",
            }
            ext = ext_map.get(content_type, ".png")
            return image_bytes, ext, content_type
    return None


def upload_image_bytes(client, image_bytes, filename, mime_type):
    """Upload image bytes naar AskDelphi /resource endpoint."""
    with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as tmp:
        tmp.write(image_bytes)
        tmp_path = tmp.name

    try:
        with open(tmp_path, "rb") as f:
            files = {"File": (filename, f, mime_type)}
            response = api.upload_resource(client, files)
        return response
    finally:
        Path(tmp_path).unlink(missing_ok=True)


def build_image_embed_html(client, resource_response):
    """Bouw een <doppio-embedded> tag van een upload_image response."""
    res = resource_response["resource"]
    topic_guid = res["topicGuid"]
    title = res["title"]
    thumbnail = res["thumbnailImageBase64"]

    link = (
        f"tenant/{client.tenant_id}"
        f"/project/{client.project_id}"
        f"/acl/{client.acl_entry_id}"
        f"/topic/{topic_guid}/edit"
    )

    return (
        f'<p><doppio-embedded '
        f'target="{topic_guid}" '
        f'use="default" view="default" '
        f'title="{title}" '
        f'thumbnail="data:image/png;base64,{thumbnail}" '
        f'link="{link}">'
        f'\xa0</doppio-embedded></p>'
    )


def cell_to_html(cell, client=None, document=None):
    """Convert a docx table cell to an HTML string."""
    html_parts = []

    for paragraph in cell.paragraphs:
        pPr = paragraph._p.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr")) if pPr is not None else None

        runs_html = []
        for run in paragraph.runs:
            if client is not None:
                image_data = extract_image_from_run(run, document)
                if image_data is not None:
                    image_bytes, ext, content_type = image_data
                    filename = f"image{ext}"
                    resource = upload_image_bytes(client, image_bytes, filename, content_type)
                    embed_html = build_image_embed_html(client, resource)
                    runs_html.append(embed_html)
                    continue

            text = run.text
            if not text:
                continue
            text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            runs_html.append(text)

        line = "".join(runs_html)
        if not line:
            continue

        if numPr is not None:
            ilvl_el = numPr.find(qn("w:ilvl"))
            level = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0
            html_parts.append(("li", level, line))
        else:
            html_parts.append(("p", None, line))

    output = []
    current_level = -1

    for tag, level, content in html_parts:
        if tag == "li":
            while current_level < level:
                output.append("<ul>")
                current_level += 1
            while current_level > level:
                output.append("</ul>")
                current_level -= 1
            output.append(f"<li>{content}</li>")
        else:
            while current_level >= 0:
                output.append("</ul>")
                current_level -= 1
            output.append(f"<p>{content}</p>")

    while current_level >= 0:
        output.append("</ul>")
        current_level -= 1

    return "\n".join(output)


def convert_doc_to_tables(path, client=None):
    """Return (text_tables, html_tables): two parallel lists of DataFrames."""
    from docx import Document

    document = Document(path)

    text_tables = []
    html_tables = []

    for table in document.tables:
        text_data = []
        html_data = []

        for row in table.rows:
            text_row = []
            html_row = []
            for cell in row.cells:
                text_row.append(cell.text.strip())
                html_row.append(cell_to_html(cell, client=client, document=document))
            text_data.append(text_row)
            html_data.append(html_row)

        if text_data:
            text_df = pd.DataFrame(text_data[1:], columns=text_data[0])
            html_df = pd.DataFrame(html_data[1:], columns=text_data[0])
            text_tables.append(text_df)
            html_tables.append(html_df)

    print(f"retrieved {len(text_tables)} tables from doc {path}")
    return text_tables, html_tables


def filter_tables_by_title(table_list, keyword):
    return [table for table in table_list if table.columns[0].lower().startswith(keyword)]


def filter_tables_pair_by_title(text_tables, html_tables, keyword):
    """Filter both table lists in lockstep, returning matched pairs."""
    pairs = [
        (tt, ht)
        for tt, ht in zip(text_tables, html_tables)
        if tt.columns[0].lower().startswith(keyword)
    ]
    if not pairs:
        return [], []
    return [p[0] for p in pairs], [p[1] for p in pairs]


def remove_prefix_ci(text, prefix):
    if text.lower().startswith(prefix.lower()):
        return text[len(prefix):]
    return text


def clean_strip(text):
    result = re.sub(r"^[^a-zA-Z]+", "", text)
    result = re.sub(r"[^a-zA-Z]+$", "", result)
    return result


def extract_digicoach_title(text_tables):
    title_tables = filter_tables_by_title(text_tables, "titel")
    if len(title_tables) == 0:
        raise ValueError("No title tables found")
    if len(title_tables) > 1:
        warnings.warn("Multiple title tables found, continuing with the first", UserWarning)
    title = title_tables[0].columns[1]
    print(f"Found title: {title}")
    return title


def extract_digicoach_tags(text_tables):
    tags = []
    tag_tables = filter_tables_by_title(text_tables, "tag")
    if len(tag_tables) == 0:
        raise ValueError("No tag tables found")
    if len(tag_tables) > 1:
        warnings.warn("Multiple tag tables found, continuing with the first", UserWarning)
    tag_table = tag_tables[0]
    for i in range(0, len(tag_table.columns), 2):
        sub_table = tag_table.iloc[:, i : i + 2]
        hits = sub_table.replace("", np.nan).dropna()
        values = hits.iloc[:, 0].to_list()
        if values:
            new_tag = {}
            tag_type = sub_table.columns[0]
            tag_type = remove_prefix_ci(tag_type, "tag")
            tag_type = clean_strip(tag_type)
            new_tag["type"] = tag_type
            new_tag["values"] = values
            for val in values:
                print(f"Found tag: {tag_type} {val}")
            tags.append(new_tag)
    return tags


def extract_digicoach_tasks(text_tables, html_tables):
    """Extract tasks with HTML descriptions."""
    tasks = []
    text_task_tables, html_task_tables = filter_tables_pair_by_title(
        text_tables, html_tables, "taak"
    )
    if not text_task_tables:
        raise ValueError("No task tables found")

    for text_table, html_table in zip(text_task_tables, html_task_tables):
        new_task = {}

        text_data = text_table.iloc[:, 1].replace("", np.nan).dropna()
        html_data = html_table.iloc[:, 1]

        task_name = text_data.name
        task_name = remove_prefix_ci(task_name, "taak")
        task_name = clean_strip(task_name)
        new_task["name"] = task_name

        new_task["description"] = html_data.iloc[text_data.index[0]]

        if not len(text_data) % 2 == 1:
            raise ValueError(f"invalid step_df: {text_data}")

        step_indices = text_data.index[1:]
        steps = []
        for i in range(0, len(step_indices), 2):
            idx_name = step_indices[i]
            idx_desc = step_indices[i + 1]

            step_name = text_data[idx_name]
            step_name = remove_prefix_ci(step_name, "stap")
            step_name = clean_strip(step_name)

            new_step = {
                "name": step_name,
                "description": html_data.iloc[idx_desc],
            }
            steps.append(new_step)
            print(f"Found step: {step_name}")

        new_task["steps"] = steps
        tasks.append(new_task)
        print(f"Found task: {task_name}")

    return tasks


def extract_digicoach_sources(text_tables):
    sources = []
    source_tables = filter_tables_by_title(text_tables, "nr")
    if len(source_tables) == 0:
        raise ValueError("No source tables found")

    for table in source_tables:
        data = table.iloc[:, [1, 2, 3]]
        data = data.replace("", np.nan).dropna()

        instruction_df = pd.DataFrame(
            data.values.reshape(-1, 3), columns=["Titel", "Type", "Link"]
        )

        for _, row in instruction_df.iterrows():
            new_source = {}
            source_titel = clean_strip(row["Titel"])
            new_source["titel"] = source_titel
            new_source["type"] = row["Type"]
            new_source["link"] = row["Link"]
            sources.append(new_source)

        print(f"Found sources: {sources}")

    return sources


def convert_doc_to_json(path, client=None):
    text_tables, html_tables = convert_doc_to_tables(path, client=client)
    title = extract_digicoach_title(text_tables)
    tags = extract_digicoach_tags(text_tables)
    tasks = extract_digicoach_tasks(text_tables, html_tables)
    sources = extract_digicoach_sources(text_tables)
    digicoach = {"name": title, "tags": tags, "tasks": tasks, "sources": sources}
    digicoach_json = json.dumps(digicoach, indent=2)
    return digicoach_json
